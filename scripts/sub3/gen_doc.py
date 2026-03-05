from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "C:/Users/Aditya/OneDrive/Desktop/CF/campus-flow/CampusFlow_API_Changes.docx"

doc = Document()

def set_font(run, bold=False, size=11, color=None, mono=False):
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x06, 0x4E, 0x3B)
    return p

def heading3(text):
    return doc.add_heading(text, level=3)

def para(text="", bold=False, mono=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def endpoint_row(method, path, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(" " + method + " ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.name = "Courier New"
    colors = {"POST": (0x16, 0xA3, 0x4A), "GET": (0x25, 0x63, 0xEB)}
    rgb = colors.get(method, (0x6B, 0x72, 0x80))
    r1.font.color.rgb = RGBColor(255, 255, 255)
    rPr = r1._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*rgb))
    rPr.append(shd)
    r2 = p.add_run("  " + path)
    r2.bold = True
    r2.font.name = "Courier New"
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    r3 = p.add_run("   -   " + description)
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Title ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("CampusFlow")
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Backend API Changes -- Developer Handoff")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run(datetime.date.today().strftime("%B %d, %Y"))
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_page_break()

# ── 1. Overview ────────────────────────────────────────────────────────────────
heading1("1. Overview")
para(
    "This document describes all new backend endpoints, models, and schemas added as part of "
    "the latest feature sprint for the CampusFlow admin and vendor dashboards. "
    "No existing endpoints were modified -- all changes are purely additive."
)
doc.add_paragraph()
bullet("4 new features implemented across Admin and Vendor roles")
bullet("8 new API endpoints total")
bullet("1 new database table: admin_deduct_requests")
bullet("3 new Pydantic request schemas")
divider()

# ── 2. New Database Model ──────────────────────────────────────────────────────
heading1("2. New Database Model")
heading2("AdminDeductRequest")
para(
    "Stores deduct requests submitted by vendors. Admin reviews and approves or rejects them. "
    "Table name: admin_deduct_requests"
)
code_block(
    "class AdminDeductRequest(Base):\n"
    "    __tablename__ = 'admin_deduct_requests'\n\n"
    "    id          = Column(UUID, primary_key=True, default=uuid.uuid4)\n"
    "    vendor_id   = Column(UUID, ForeignKey('users.id'))\n"
    "    student_id  = Column(UUID, ForeignKey('users.id'))\n"
    "    amount      = Column(Float, nullable=False)\n"
    "    reason      = Column(String, nullable=False)\n"
    "    status      = Column(String, default='PENDING')  # PENDING | APPROVED | REJECTED\n"
    "    created_at  = Column(DateTime, default=datetime.utcnow)\n"
    "    resolved_at = Column(DateTime, nullable=True)\n"
    "    resolved_by = Column(UUID, ForeignKey('users.id'), nullable=True)\n\n"
    "    vendor   = relationship('User', foreign_keys=[vendor_id])\n"
    "    student  = relationship('User', foreign_keys=[student_id])\n"
    "    resolver = relationship('User', foreign_keys=[resolved_by])"
)
para(
    "IMPORTANT: This table is new and must be created via migration (or recreate the DB in dev). "
    "Run alembic autogenerate or ensure Base.metadata.create_all() is called on startup."
)
p = doc.add_paragraph()
r = p.add_run("Status lifecycle: "); r.bold = True
p.add_run("PENDING -> APPROVED or REJECTED. Once resolved, the record is immutable.")
divider()

# ── 3. Schemas ─────────────────────────────────────────────────────────────────
heading1("3. New Pydantic Schemas")
para("File: backend/app/schemas/__init__.py", mono=True, size=10)
doc.add_paragraph()

heading3("AdminUserCreate")
para("Used by POST /admin/create-user to create any user type.")
code_block(
    "class AdminUserCreate(BaseModel):\n"
    "    username:      str\n"
    "    email:         str\n"
    "    password:      str\n"
    "    role:          str            # 'STUDENT' | 'VENDOR' | 'ADMIN'\n"
    "    student_id:    Optional[str]  # required when role == STUDENT\n"
    "    vendor_code:   Optional[str]  # required when role == VENDOR\n"
    "    business_name: Optional[str]  # required when role == VENDOR"
)

heading3("ManualDeductRequest")
para("Used by POST /admin/manual-deduct.")
code_block(
    "class ManualDeductRequest(BaseModel):\n"
    "    user_identifier: str    # email, student_id, or name\n"
    "    amount:          float\n"
    "    reason:          str"
)

heading3("AdminDeductRequestCreate")
para("Used by POST /vendor/request-admin-deduct.")
code_block(
    "class AdminDeductRequestCreate(BaseModel):\n"
    "    student_identifier: str   # email, student_id, or name\n"
    "    amount:             float\n"
    "    reason:             str"
)
divider()

# ── 4. Admin Endpoints ─────────────────────────────────────────────────────────
heading1("4. New Admin Endpoints")
para("All endpoints require:  Authorization: Bearer <admin_jwt_token>", mono=True, size=10)
doc.add_paragraph()

heading2("4.1  Create User")
endpoint_row("POST", "/admin/create-user", "Create a STUDENT, VENDOR, or ADMIN account")
doc.add_paragraph()
para("Request body (JSON):", bold=True)
code_block(
    '{\n'
    '  "username":      "Alice Smith",\n'
    '  "email":         "alice@campus.edu",\n'
    '  "password":      "securepass123",\n'
    '  "role":          "STUDENT",\n'
    '  "student_id":    "CS-2024-042",      // only for STUDENT\n'
    '  "vendor_code":   "VND-001",          // only for VENDOR\n'
    '  "business_name": "Campus Canteen"    // only for VENDOR\n'
    '}'
)
para("Success response (201):", bold=True)
code_block('{ "status": "SUCCESS", "message": "User created successfully", "user_id": "<uuid>" }')
para("Notes:", bold=True)
bullet("For VENDOR role, a vendor_profiles row is also created automatically.")
bullet("Duplicate email or student_id/vendor_code returns 400.")
bullet("role is case-insensitive.")
doc.add_paragraph()

heading2("4.2  Manual Wallet Deduction")
endpoint_row("POST", "/admin/manual-deduct", "Deduct an amount from any user's wallet")
doc.add_paragraph()
para("Request body (JSON):", bold=True)
code_block(
    '{\n'
    '  "user_identifier": "CS-2024-042",  // email, student_id, or name\n'
    '  "amount":          250.00,\n'
    '  "reason":          "Library fine"\n'
    '}'
)
para("Success response (200):", bold=True)
code_block(
    '{\n'
    '  "status":         "SUCCESS",\n'
    '  "message":        "Deducted from Alice Smith",\n'
    '  "new_balance":    750.00,\n'
    '  "transaction_id": "txn-<uuid>"\n'
    '}'
)
para("Notes:", bold=True)
bullet("Creates a Transaction with type='FEE', status='COMPLETED', receiver_id=NULL.")
bullet("Returns 400 if amount <= 0 or insufficient balance.")
bullet("Returns 404 if user_identifier matches no user.")
doc.add_paragraph()

heading2("4.3  Student Profile Lookup")
endpoint_row("GET", "/admin/student/{identifier}", "Full profile of a student")
doc.add_paragraph()
para("Path parameter:", bold=True)
bullet("{identifier} -- student_id (e.g. CS-2024-042), email, or name")
para("Success response (200):", bold=True)
code_block(
    '{\n'
    '  "user": { "id": "<uuid>", "name": "Alice", "email": "...",\n'
    '             "student_id": "CS-2024-042", "wallet_balance": 750.0,\n'
    '             "created_at": "2025-01-15T10:30:00" },\n'
    '  "transactions": [\n'
    '    { "transaction_id": "txn-...", "type": "FEE",\n'
    '      "amount": -250.0, "status": "COMPLETED", "timestamp": "..." }\n'
    '  ],\n'
    '  "fines": [\n'
    '    { "fine_id": "fine-...", "amount": 50.0, "reason": "Late return",\n'
    '      "status": "UNPAID", "issued_at": "..." }\n'
    '  ],\n'
    '  "subscriptions": [\n'
    '    { "subscription_id": "sub-...", "plan_name": "Meal Plan",\n'
    '      "amount": 500.0, "billing_cycle": "MONTHLY",\n'
    '      "next_billing_date": "2025-02-01" }\n'
    '  ]\n'
    '}'
)
para("Notes:", bold=True)
bullet("Transaction amounts are signed: negative = outflow from student, positive = inflow.")
bullet("Returns up to 20 most recent transactions, newest first.")
bullet("Only active subscriptions (is_active=True) are returned.")
doc.add_paragraph()

heading2("4.4  List Deduct Requests")
endpoint_row("GET", "/admin/deduct-requests", "List all vendor deduct requests (all statuses)")
doc.add_paragraph()
para("Success response (200) -- array:", bold=True)
code_block(
    '[\n'
    '  {\n'
    '    "id":                 "<uuid>",\n'
    '    "vendor_id":          "<uuid>",\n'
    '    "vendor_name":        "Campus Canteen",\n'
    '    "student_id":         "<uuid>",\n'
    '    "student_name":       "Alice Smith",\n'
    '    "student_identifier": "CS-2024-042",\n'
    '    "amount":             150.0,\n'
    '    "reason":             "Monthly meal dues",\n'
    '    "status":             "PENDING",\n'
    '    "created_at":         "2025-01-20T09:00:00",\n'
    '    "resolved_at":        null\n'
    '  }\n'
    ']'
)
doc.add_paragraph()

heading2("4.5  Approve Deduct Request")
endpoint_row("POST", "/admin/deduct-requests/{request_id}/approve", "Approve and execute transfer")
doc.add_paragraph()
para("What happens on approval:", bold=True)
bullet("student.wallet_balance -= amount")
bullet("vendor_user.wallet_balance += amount")
bullet("Request status -> APPROVED, resolved_at and resolved_by are populated")
bullet("Transaction created: type='VENDOR_PAYMENT', sender=student, receiver=vendor")
para("Success response (200):", bold=True)
code_block(
    '{\n'
    '  "status":         "SUCCESS",\n'
    '  "message":        "Deduct request approved and amount transferred",\n'
    '  "transaction_id": "txn-<uuid>"\n'
    '}'
)
bullet("Returns 400 if request is not PENDING or student has insufficient balance.")
doc.add_paragraph()

heading2("4.6  Reject Deduct Request")
endpoint_row("POST", "/admin/deduct-requests/{request_id}/reject", "Reject without any transfer")
doc.add_paragraph()
para("No money moves. Status is set to REJECTED, resolved_at and resolved_by are populated.")
para("Success response (200):", bold=True)
code_block('{ "status": "SUCCESS", "message": "Deduct request rejected" }')
divider()

# ── 5. Vendor Endpoints ────────────────────────────────────────────────────────
heading1("5. New Vendor Endpoints")
para("All endpoints require:  Authorization: Bearer <vendor_jwt_token>", mono=True, size=10)
doc.add_paragraph()

heading2("5.1  Submit Deduct Request to Admin")
endpoint_row("POST", "/vendor/request-admin-deduct", "Ask admin to deduct student wallet for dues")
doc.add_paragraph()
para("Request body (JSON):", bold=True)
code_block(
    '{\n'
    '  "student_identifier": "CS-2024-042",  // email, student_id, or name\n'
    '  "amount":             150.00,\n'
    '  "reason":             "Monthly meal dues Jan 2025"\n'
    '}'
)
para("Success response (200):", bold=True)
code_block(
    '{\n'
    '  "status":     "SUCCESS",\n'
    '  "request_id": "<uuid>",\n'
    '  "message":    "Deduct request sent to admin for approval"\n'
    '}'
)
para("Notes:", bold=True)
bullet("Creates a new AdminDeductRequest with status=PENDING. No money moves yet.")
bullet("Returns 404 if student_identifier matches no student.")
bullet("Returns 400 if amount <= 0.")
doc.add_paragraph()

heading2("5.2  List Own Deduct Requests")
endpoint_row("GET", "/vendor/deduct-requests", "List this vendor's submitted deduct requests")
doc.add_paragraph()
para("Success response (200) -- array:", bold=True)
code_block(
    '[\n'
    '  {\n'
    '    "id":                 "<uuid>",\n'
    '    "student_name":       "Alice Smith",\n'
    '    "student_identifier": "CS-2024-042",\n'
    '    "amount":             150.0,\n'
    '    "reason":             "Monthly meal dues Jan 2025",\n'
    '    "status":             "PENDING",\n'
    '    "created_at":         "2025-01-20T09:00:00",\n'
    '    "resolved_at":        null\n'
    '  }\n'
    ']'
)
divider()

# ── 6. Auth ────────────────────────────────────────────────────────────────────
heading1("6. Authentication")
para("All new endpoints use Bearer JWT auth -- same as existing endpoints. No changes to auth flow.")
code_block("Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...")
bullet("Admin endpoints reject non-ADMIN tokens with 403.")
bullet("Vendor endpoints reject non-VENDOR tokens with 403.")
bullet("Token expiry: 300 minutes (unchanged).")
divider()

# ── 7. Error Reference ─────────────────────────────────────────────────────────
heading1("7. Error Reference")
table = doc.add_table(rows=1, cols=3)
table.style = "Light List Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "HTTP Status"
hdr[1].text = "Error detail"
hdr[2].text = "Endpoint(s)"
rows_data = [
    ("400", "Invalid role",                           "POST /admin/create-user"),
    ("400", "Email already registered",               "POST /admin/create-user"),
    ("400", "Student ID already registered",          "POST /admin/create-user"),
    ("400", "vendor_code and business_name required", "POST /admin/create-user"),
    ("400", "Vendor code already exists",             "POST /admin/create-user"),
    ("400", "Amount must be positive",                "POST /admin/manual-deduct, /vendor/request-admin-deduct"),
    ("400", "Insufficient balance",                   "POST /admin/manual-deduct"),
    ("400", "Student has insufficient balance",       "POST /admin/deduct-requests/{id}/approve"),
    ("400", "Request already resolved",               "POST /admin/deduct-requests/{id}/approve|reject"),
    ("404", "User not found",                         "POST /admin/manual-deduct"),
    ("404", "Student not found",                      "GET /admin/student/{id}, /vendor/request-admin-deduct"),
    ("404", "Request not found",                      "POST /admin/deduct-requests/{id}/approve|reject"),
]
for s, d, e in rows_data:
    row = table.add_row().cells
    row[0].text = s; row[1].text = d; row[2].text = e
doc.add_paragraph()
divider()

# ── 8. Transaction Types ───────────────────────────────────────────────────────
heading1("8. Transaction Type Reference")
table2 = doc.add_table(rows=1, cols=2)
table2.style = "Light List Accent 1"
h = table2.rows[0].cells
h[0].text = "type"
h[1].text = "Description"
types = [
    ("P2P",            "Student-to-student transfer"),
    ("VENDOR_PAYMENT", "Payment to a vendor (direct or via deduct request approval)"),
    ("SUB",            "Subscription billing"),
    ("FINE",           "Fine issued by admin (money leaves system)"),
    ("FEE",            "Manual deduction by admin (money leaves system)"),
    ("TOP_UP",         "Wallet top-up"),
]
for t, d in types:
    r = table2.add_row().cells
    r[0].text = t; r[1].text = d
doc.add_paragraph()
divider()

# ── 9. Migration ───────────────────────────────────────────────────────────────
heading1("9. Migration / Setup Note")
para("The admin_deduct_requests table is brand new and must be created before starting the server.", bold=True)
bullet("Dev (SQLite): Ensure Base.metadata.create_all(bind=engine) is called in main.py on startup.")
bullet(
    "Prod (PostgreSQL): Run alembic revision --autogenerate -m 'add_admin_deduct_requests' "
    "then alembic upgrade head"
)
para("No existing tables were altered. All other migrations remain valid.")
divider()

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CampusFlow  |  Invenza'26 Hackathon  |  Confidential")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.save(OUT)
print("Done ->", OUT)
